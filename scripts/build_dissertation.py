"""
Assemble the COM4901 dissertation from per-chapter Markdown into a formatted
.docx.

WHY THIS EXISTS
---------------
The final report must be Times New Roman 12 / 1.5 spacing / 1 inch margins with
IEEE numeric references (KIU guidelines section 10.2). Drafting directly in Word
makes revisions unreviewable -- no diffs, no version control, and the formatting
has to be re-applied by hand every pass. So chapters are written as Markdown in
created_docs/dissertation/ and this script applies the formatting once,
mechanically, on every build.

CITATIONS
---------
Chapters cite by STABLE KEY -- [@peffers2007] -- never by number. This script
resolves keys to IEEE numbers in order of first appearance across the assembled
document and emits the matching numbered reference list. That means P9 can take
the bibliography from 29 entries to 40+, in any order, WITHOUT touching a single
in-text citation. A key with no entry in references.md is a hard error: a
dissertation that cites [17] where no [17] exists is worse than one that fails
to build.

WHAT IT SUPPORTS
----------------
Headings (#, ##, ###), paragraphs, inline **bold** / *italic* / `code`, bullet
and numbered lists, pipe tables, figures (![caption](path)) and block quotes.
That is the subset the chapter drafts actually use; anything else is passed
through as plain text rather than silently mangled.

Run:  python scripts/build_dissertation.py [--out created_docs/Final_Report.docx]
Needs: python-docx (installed).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    import docx
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit('ERROR: python-docx is required.  pip install python-docx')

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / 'created_docs' / 'dissertation'
REFS_FILE = SRC / 'references.md'

FONT = 'Times New Roman'
SIZE = Pt(12)
LINE_SPACING = 1.5

TITLE_PAGE = {
    'title': 'A Cost-Effective Predictive HR Analytics Framework '
             'for Sri Lankan SMEs Using Cloud-Native Serverless AI',
    'student': 'Theekshana Gimhan',
    'student_id': '15002',
    'supervisor': 'Ms. Thanuja Irugalbandara',
    'module': 'COM 4901 - Final Year Individual Project',
    'degree': 'BSc (Hons) in Management Information Systems',
    'faculty': 'Faculty of Computer Science and Engineering',
    'university': 'KIU University',
    'date': 'August 2026',
}

# Assembly order. A chapter absent from disk is skipped with a notice rather
# than failing the build, so partial drafts still produce a readable document.

# Front matter, rendered between the title page and the table of contents.
# Deliberately excluded from the body word count: the 10,000-word minimum is
# counted on the chapters, so keeping the abstract out of it keeps the reported
# figure conservative rather than flattering.
FRONT_MATTER = [
    'abstract.md',
    'acknowledgements.md',
]

CHAPTERS = [
    'ch1_introduction.md',
    'ch2_literature_review.md',
    'ch3_methodology.md',
    'ch4_implementation.md',
    'ch5_evaluation.md',
    'ch6_conclusions.md',
]

# Back matter, rendered after the reference list. Also excluded from the body
# word count -- appendices are supporting material, not argument, and counting
# them towards the 10,000 minimum would flatter the figure.
BACK_MATTER = [
    'appendices.md',
]


# ---------------------------------------------------------------------------
# References
# ---------------------------------------------------------------------------
def load_references() -> dict:
    """Parse references.md -> {key: entry}.

    Entries live in pipe tables whose first cell is a backticked key. Header
    and separator rows have no backticked first cell and are skipped.
    """
    if not REFS_FILE.exists():
        sys.exit('ERROR: missing %s' % REFS_FILE)
    refs = {}
    for line in REFS_FILE.read_text(encoding='utf8').splitlines():
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        if len(cells) < 2:
            continue
        m = re.fullmatch(r'`([A-Za-z0-9]+)`', cells[0])
        if not m:
            continue
        refs[m.group(1)] = cells[1]
    if not refs:
        sys.exit('ERROR: no reference entries parsed from references.md')
    return refs


class Citations:
    """Assigns IEEE numbers in order of first appearance."""

    def __init__(self, refs: dict):
        self.refs = refs
        self.order = []          # keys, in citation order
        self.unresolved = set()

    def number(self, key: str) -> int:
        if key not in self.refs:
            self.unresolved.add(key)
            return 0
        if key not in self.order:
            self.order.append(key)
        return self.order.index(key) + 1

    def substitute(self, text: str) -> str:
        def repl(m):
            n = self.number(m.group(1))
            return '[%d]' % n if n else '[??%s]' % m.group(1)
        return re.sub(r'\[@([A-Za-z0-9]+)\]', repl, text)


# ---------------------------------------------------------------------------
# Document styling
# ---------------------------------------------------------------------------
def configure(doc: Document) -> None:
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = SIZE
    # East-Asian font mapping, else Word substitutes for some glyphs.
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)

    for name, size, bold in (('Heading 1', 16, True),
                             ('Heading 2', 14, True),
                             ('Heading 3', 12, True)):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)   # Word's default heading blue is wrong here
        st.paragraph_format.line_spacing = LINE_SPACING
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_page_numbers(doc: Document) -> None:
    """Footer field: PAGE. Word renders it; python-docx cannot compute it."""
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for instr, kind in (('begin', 'w:fldCharType'),):
            el = OxmlElement('w:fldChar')
            el.set(qn(kind), instr)
            run._r.append(el)
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = 'PAGE'
        run._r.append(it)
        el = OxmlElement('w:fldChar')
        el.set(qn('w:fldCharType'), 'end')
        run._r.append(el)
        p.style.font.name = FONT
        p.style.font.size = Pt(10)


def add_toc(doc: Document) -> None:
    """Insert a TOC field. Word populates it on open (or F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = 'Right-click here and choose "Update Field" to build the table of contents.'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------
MAX_FIG_WIDTH = 6.0    # inches -- the text column
MIN_DPI = 150          # below this, a printed screenshot looks soft


def figure_width(path) -> float:
    """Widest size that keeps the image at or above MIN_DPI, capped at the column."""
    try:
        from PIL import Image
        with Image.open(str(path)) as im:
            px = im.width
    except Exception:
        return MAX_FIG_WIDTH      # unreadable header: fall back to full width
    return max(2.5, min(MAX_FIG_WIDTH, px / float(MIN_DPI)))


# ---------------------------------------------------------------------------
# Lists of figures and tables
# ---------------------------------------------------------------------------
# KIU guidelines section 10.3 requires these as separate preliminary pages.
# They are generated from the captions in the chapter sources rather than
# maintained by hand, so renaming a figure cannot leave the list stale.
#
# No page numbers: Word computes those from fields, and a field that fails to
# resolve produces an EMPTY list, which is worse than a correct list without
# page numbers. Caption text is deterministic; page numbers here would not be.
FIG_CAPTION = re.compile(r'!\[(.*?)\]\(.+?\)')
TAB_CAPTION = re.compile(r'\*\*(Table\s+[0-9A-Z]+\.[0-9]+\s+.*?)\*\*')


def collect_captions(names: list) -> tuple:
    """Scan sources in assembly order -> (figure captions, table captions)."""
    figures, tables = [], []
    for name in names:
        path = SRC / name
        if not path.exists():
            continue
        for line in path.read_text(encoding='utf8').splitlines():
            stripped = line.strip()
            m = FIG_CAPTION.fullmatch(stripped)
            if m:
                figures.append(strip_inline(m.group(1)))
                continue
            m = TAB_CAPTION.fullmatch(stripped)
            if m:
                tables.append(strip_inline(m.group(1)))
    return figures, tables


def add_caption_list(doc: Document, heading: str, captions: list) -> None:
    doc.add_heading(heading, level=1)
    if not captions:
        doc.add_paragraph('None.')
        return
    for caption in captions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.add_run(caption)


# ---------------------------------------------------------------------------
# Inline markdown
# ---------------------------------------------------------------------------
INLINE = re.compile(r'(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)')


def add_runs(par, text: str) -> None:
    """Render **bold**, *italic* and `code` as real runs."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith('**') and piece.endswith('**'):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith('`') and piece.endswith('`'):
            r = par.add_run(piece[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10.5)
        elif piece.startswith('*') and piece.endswith('*'):
            par.add_run(piece[1:-1]).italic = True
        else:
            par.add_run(piece)


def strip_inline(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


# ---------------------------------------------------------------------------
# Block markdown
# ---------------------------------------------------------------------------
def is_table_sep(line: str) -> bool:
    return bool(re.fullmatch(r'\|[\s:|-]+\|', line.strip()))


def render(doc: Document, md: str, cites: Citations, counter: dict) -> None:
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = cites.substitute(raw.rstrip())
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Table: a pipe row followed by a separator row.
        if (stripped.startswith('|') and i + 1 < len(lines)
                and is_table_sep(lines[i + 1])):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            body = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = cites.substitute(lines[i].strip())
                body.append([c.strip() for c in row.strip('|').split('|')])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Table Grid'
            for cell, txt in zip(table.rows[0].cells, header):
                cell.paragraphs[0].clear() if hasattr(cell.paragraphs[0], 'clear') else None
                r = cell.paragraphs[0].add_run(strip_inline(txt))
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(10.5)
            for row in body:
                cells = table.add_row().cells
                for cell, txt in zip(cells, row):
                    par = cell.paragraphs[0]
                    add_runs(par, txt)
                    for r in par.runs:
                        if r.font.name != 'Consolas':
                            r.font.name = FONT
                        r.font.size = Pt(10.5)
            doc.add_paragraph()
            continue

        # Figure
        m = re.fullmatch(r'!\[(.*?)\]\((.+?)\)', stripped)
        if m:
            caption, path = m.group(1), (ROOT / m.group(2))
            if path.exists():
                # Never upscale past MIN_DPI. A 6in-wide frame makes a small
                # screenshot soft in print; better a narrower sharp figure than
                # a full-width blurry one.
                doc.add_picture(str(path), width=Inches(figure_width(path)))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph('[missing figure: %s]' % m.group(2))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(10.5)
            i += 1
            continue

        if stripped.startswith('### '):
            doc.add_heading(strip_inline(stripped[4:]), level=3)
        elif stripped.startswith('## '):
            doc.add_heading(strip_inline(stripped[3:]), level=2)
        elif stripped.startswith('# '):
            doc.add_page_break() if counter['chapters'] else None
            counter['chapters'] += 1
            doc.add_heading(strip_inline(stripped[2:]), level=1)
        elif stripped.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote'
                                  if 'Intense Quote' in [s.name for s in doc.styles]
                                  else None)
            add_runs(p, stripped[2:])
            counter['words'] += len(strip_inline(stripped[2:]).split())
        elif re.match(r'^[-*] ', stripped):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, stripped[2:])
            counter['words'] += len(strip_inline(stripped[2:]).split())
        elif re.match(r'^\d+\. ', stripped):
            p = doc.add_paragraph(style='List Number')
            add_runs(p, re.sub(r'^\d+\.\s*', '', stripped))
            counter['words'] += len(strip_inline(stripped).split())
        elif set(stripped) == {'-'} and len(stripped) >= 3:
            pass  # horizontal rule -- no Word equivalent worth emitting
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, stripped)
            counter['words'] += len(strip_inline(stripped).split())
        i += 1


def build_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    for text, size, bold in (
            (TITLE_PAGE['title'], 18, True),
            ('', 12, False),
            (TITLE_PAGE['module'], 13, False),
            (TITLE_PAGE['degree'], 12, False),
            ('', 12, False),
            (TITLE_PAGE['student'], 13, True),
            ('Student ID: ' + TITLE_PAGE['student_id'], 12, False),
            ('', 12, False),
            ('Supervisor: ' + TITLE_PAGE['supervisor'], 12, False),
            ('', 12, False),
            (TITLE_PAGE['faculty'], 12, False),
            (TITLE_PAGE['university'], 12, False),
            (TITLE_PAGE['date'], 12, False)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = FONT
    doc.add_page_break()


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument('--out', default=str(ROOT / 'created_docs' / 'Final_Report_COM4901_Theekshana_Gimhan.docx'))
    args = ap.parse_args()

    refs = load_references()
    cites = Citations(refs)
    counter = {'words': 0, 'chapters': 0}

    doc = Document()
    configure(doc)
    add_page_numbers(doc)
    build_title_page(doc)

    print('=' * 74)
    print('BUILDING DISSERTATION')
    print('=' * 74)

    # Front matter first, counted separately so it stays out of the body total.
    front = {'words': 0, 'chapters': 0}
    for name in FRONT_MATTER:
        path = SRC / name
        if not path.exists():
            print('  skip   %-28s (not drafted yet)' % name)
            continue
        before = front['words']
        render(doc, path.read_text(encoding='utf8'), cites, front)
        print('  ok     %-28s %5d words (front matter)' % (name, front['words'] - before))
        doc.add_page_break()

    doc.add_heading('Table of Contents', level=1)
    add_toc(doc)
    doc.add_page_break()

    figures, tables = collect_captions(CHAPTERS + BACK_MATTER)
    add_caption_list(doc, 'List of Figures', figures)
    doc.add_page_break()
    add_caption_list(doc, 'List of Tables', tables)
    doc.add_page_break()
    print('  ok     %-28s %5d figures, %d tables'
          % ('lists of figures/tables', len(figures), len(tables)))

    found = 0
    for name in CHAPTERS:
        path = SRC / name
        if not path.exists():
            print('  skip   %-28s (not drafted yet)' % name)
            continue
        before = counter['words']
        render(doc, path.read_text(encoding='utf8'), cites, counter)
        print('  ok     %-28s %5d words' % (name, counter['words'] - before))
        found += 1

    if not found:
        sys.exit('ERROR: no chapter files found in %s' % SRC)

    # Back matter cites too, but is rendered AFTER the reference list. Register
    # its keys first so those references exist and are numbered in reading
    # order; substitute() is idempotent, so re-running it during the render
    # below yields the same numbers.
    for name in BACK_MATTER:
        path = SRC / name
        if path.exists():
            cites.substitute(path.read_text(encoding='utf8'))

    # References, numbered in order of first citation.
    doc.add_page_break()
    doc.add_heading('References', level=1)
    for idx, key in enumerate(cites.order, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        add_runs(p, '[%d] %s' % (idx, refs[key]))

    back = {'words': 0, 'chapters': 1}   # chapters=1 -> page break before it
    for name in BACK_MATTER:
        path = SRC / name
        if not path.exists():
            print('  skip   %-28s (not drafted yet)' % name)
            continue
        render(doc, path.read_text(encoding='utf8'), cites, back)
        print('  ok     %-28s %5d words (back matter)' % (name, back['words']))

    if cites.unresolved:
        print()
        print('BUILD FAILED -- unresolved citation keys:')
        for k in sorted(cites.unresolved):
            print('   [@%s]  -- add it to created_docs/dissertation/references.md' % k)
        sys.exit(1)

    out = Path(args.out)
    doc.save(out)

    print()
    print('-' * 74)
    print('Body word count (excl. title page, ToC, references) : %s'
          % format(counter['words'], ','))
    print('Minimum required                                    : 10,000')
    remaining = 10000 - counter['words']
    print('Status                                              : %s'
          % ('MET' if remaining <= 0 else '%s words short' % format(remaining, ',')))
    print('References cited                                    : %d of %d defined'
          % (len(cites.order), len(refs)))
    uncited = [k for k in refs if k not in cites.order]
    if uncited:
        print('Defined but not yet cited                           : %s'
              % ', '.join(sorted(uncited)[:8]) + (' ...' if len(uncited) > 8 else ''))
    print('-' * 74)
    print('Wrote %s' % out)
    print()
    print('NOTE: open in Word and press Ctrl+A then F9 to populate the table of')
    print('      contents and page numbers -- Word computes those fields, not this script.')


if __name__ == '__main__':
    main()
